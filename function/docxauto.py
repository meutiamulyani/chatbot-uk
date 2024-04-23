from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from fastapi import Request
from datetime import datetime, date

class Doc_Auto():
    # doc = Document()
    def __init__(self, db_con, model):
        # Create a new Document
        self.doc = Document()
        self.db_con = db_con
        self.model = model
    
    # kondisi dokumen
    def doc_id(self, nomor_hp):
        useract = self.db_con.query(self.model.user_activity).filter_by(no_hp=nomor_hp).first()
        print(useract)

        return useract

    def wrapper_doc(self, nomor_hp):
        # panggil doc_id untuk cek query nomor telpon dan activity
        user_activity = self.doc_id(nomor_hp=nomor_hp)
        print(f'user id: {nomor_hp}')

        # MEMBUAT form
        print("lancar")
        # # data form
        reg_doc = self.db_con.query(self.model.registrasi).filter_by(id_user_activity = user_activity.id).first()
        title = self.doc.add_paragraph('REGISTRASI MEMBER BARU KOPITU\nTanggal: ............')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.bold = True
        
        # Biodata
        self.doc.add_paragraph(
            'Biodata Lengkap:'
        )
        table = self.doc.add_table(rows=6, cols=2)
        hdr_cells = table.columns[0].cells
        hdr_cells[0].text = 'Nama\t\t\t:'
        hdr_cells[1].text = 'NIK\t\t\t:'
        hdr_cells[2].text = 'Tempat/Tanggal Lahir:'
        hdr_cells[3].text = 'Jenis Kelamin\t\t:'
        hdr_cells[4].text = 'Nomor HP\t\t:'
        hdr_cells[5].text = 'Member\t\t:'
                    
        for cell in table.columns[0].cells:
            cell.width = Inches(1.75)

        cols_cells = table.columns[1].cells
        cols_cells[0].text = reg_doc.nama
        cols_cells[1].text = reg_doc.nik
        cols_cells[2].text = reg_doc.ttl
        cols_cells[3].text = reg_doc.jenis_kelamin
        cols_cells[4].text = reg_doc.no_hp
        cols_cells[5].text = reg_doc.member
        
        # Alamat
        self.doc.add_paragraph(
            'Alamat Lengkap:'
        )
        
        table_alamat = self.doc.add_table(rows=6, cols=2)
        hdr_cells = table_alamat.columns[0].cells
        hdr_cells[0].text = 'Provinsi\t\t:'
        hdr_cells[1].text = 'Kecamatan\t\t:'
        hdr_cells[2].text = 'Kabupaten/Kota\t:'
        hdr_cells[3].text = 'Kode Pos\t\t:'
        hdr_cells[4].text = 'Warganegara\t\t:'
        hdr_cells[5].text = 'Alamat\t\t:'
                    
        for cell in table_alamat.columns[0].cells:
            cell.width = Inches(1.75)

        cols_alamat = table_alamat.columns[1].cells
        cols_alamat[0].text = reg_doc.provinsi
        cols_alamat[1].text = reg_doc.kecamatan
        cols_alamat[2].text = reg_doc.kabupaten_kota
        cols_alamat[3].text = reg_doc.kode_pos
        cols_alamat[4].text = reg_doc.warganegara
        cols_alamat[5].text = reg_doc.alamat

        # Usaha Keanggotaan
        self.doc.add_paragraph(
            'Usaha Keanggotaan:'
        )
        table_member = self.doc.add_table(rows=4, cols=2)
        hdr_member = table_member.columns[0].cells
        hdr_member[0].text = 'Jenis Usaha\t\t:'
        hdr_member[1].text = 'Kelas Usaha\t\t:'
        hdr_member[2].text = 'Deskripsi Usaha\t:'
        hdr_member[3].text = 'Ijin Usaha\t\t:'
                    
        for cell in table_member.columns[0].cells:
            cell.width = Inches(1.75)

        cols_member = table_member.columns[1].cells
        cols_member[0].text = reg_doc.jenis_usaha
        cols_member[1].text = reg_doc.kelas_usaha
        cols_member[2].text = reg_doc.deskripsi_usaha
        cols_member[3].text = reg_doc.ijin_usaha

        # rename doc
        current_time = datetime.now().strftime("%d-%m-%y_%H-%M-%S")
        file_name = f"{reg_doc.nik}_{current_time}"

        today = date.today()
        uppercase_date = today.strftime("%B %d, %Y").upper()

        ttd = self.doc.add_paragraph(f'kopitu_member_{reg_doc.kabupaten_kota.upper()}, {uppercase_date}\n\n\n{reg_doc.nama}')
        ttd.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        save = self.doc.save(f'public/files/{file_name}.docx')
        
        return file_name
